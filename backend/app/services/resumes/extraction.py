"""
ISSUE-10 / ISSUE-11 / ISSUE-12 — Resume PDF/DOCX upload & text extraction,
and JD file upload & text extraction.

Resumes accept .pdf / .docx (per the upload contract). JDs additionally
accept .txt, since job descriptions are frequently just plain text. Three
independent extractors (PyMuPDF for PDF, python-docx for DOCX, a plain
decode for TXT) feed one normalized `extract_text()` entry point so the
rest of the pipeline never has to care what format the source file was.
"""
import io
import re

import docx  # python-docx
import fitz  # PyMuPDF

from app.services.resumes.exceptions import ExtractionError, UnsupportedFileTypeError

RESUME_EXTENSIONS = {".pdf", ".docx"}
JD_EXTENSIONS = {".pdf", ".docx", ".txt"}

_CONTENT_TYPE_MAP = {
    "application/pdf": ".pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "text/plain": ".txt",
}


def guess_extension(
    filename: str, content_type: str | None, allowed: set[str] = RESUME_EXTENSIONS
) -> str:
    name = (filename or "").lower().strip()
    for ext in allowed:
        if name.endswith(ext):
            return ext
    # fall back to content-type sniffing (some clients omit a filename ext)
    sniffed = _CONTENT_TYPE_MAP.get(content_type or "")
    if sniffed in allowed:
        return sniffed
    accepted = ", ".join(sorted(allowed))
    raise UnsupportedFileTypeError(
        f"Unsupported file '{filename}'. Only {accepted} are accepted."
    )


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
    except Exception as exc:  # malformed / corrupt PDF
        raise ExtractionError(f"Could not open PDF: {exc}") from exc
    try:
        pages = [page.get_text("text") for page in doc]
    finally:
        doc.close()
    text = "\n".join(pages)
    return _normalize_text(text)


def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(file_bytes))
    except Exception as exc:  # malformed docx / not a real zip
        raise ExtractionError(f"Could not open DOCX: {exc}") from exc

    parts: list[str] = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    return _normalize_text(text)


def extract_text_from_txt(file_bytes: bytes) -> str:
    try:
        text = file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = file_bytes.decode("latin-1")
        except Exception as exc:
            raise ExtractionError(f"Could not decode text file: {exc}") from exc
    return _normalize_text(text)


def _extract_by_extension(file_bytes: bytes, ext: str) -> str:
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    if ext == ".docx":
        return extract_text_from_docx(file_bytes)
    return extract_text_from_txt(file_bytes)


def extract_text(file_bytes: bytes, filename: str, content_type: str | None = None) -> str:
    """Resume text extraction entry point (ISSUE-11). Accepts .pdf/.docx."""
    ext = guess_extension(filename, content_type, allowed=RESUME_EXTENSIONS)
    text = _extract_by_extension(file_bytes, ext)

    if not text.strip():
        raise ExtractionError(
            "No extractable text found in the uploaded resume "
            "(it may be a scanned/image-only document)."
        )
    return text


def extract_jd_file(file_bytes: bytes, filename: str, content_type: str | None = None) -> str:
    """JD text extraction entry point (ISSUE-12). Accepts .pdf/.docx/.txt —
    JDs are commonly shared as plain text as well as formatted documents."""
    ext = guess_extension(filename, content_type, allowed=JD_EXTENSIONS)
    text = _extract_by_extension(file_bytes, ext)

    if not text.strip():
        raise ExtractionError(
            "No extractable text found in the uploaded job description "
            "(it may be a scanned/image-only document)."
        )
    return text


def extract_jd_text(text: str) -> str:
    """Normalize pasted job-description text using the same validation."""
    normalized = _normalize_text(text)
    if not normalized:
        raise ExtractionError("Job description text is empty.")
    return normalized


def _normalize_text(text: str) -> str:
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # collapse runs of blank lines/spaces produced by PDF text extraction
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
